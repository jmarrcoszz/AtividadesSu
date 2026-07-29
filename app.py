import streamlit as st
import io
import os
import re
import tempfile
from datetime import datetime
from copy import deepcopy
from docx import Document
from docx.document import Document as DocxDocument
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph


NOME_TEMPLATE = "template_atividade.docx"
OPCOES_ANO = ["3º ANO", "4º ANO", "5º ANO"]


def inicializar_estado():
    """Inicializa st.session_state com valores padrão."""
    if "questoes" not in st.session_state:
        st.session_state.questoes = []
    if "professor" not in st.session_state:
        st.session_state.professor = "Susanne"
    if "componente" not in st.session_state:
        st.session_state.componente = ""
    if "data_atividade" not in st.session_state:
        st.session_state.data_atividade = ""
    if "ano_serie" not in st.session_state:
        st.session_state.ano_serie = "4º ANO"


def obter_caminho_template():
    """Retorna o caminho absoluto do template na raiz do projeto."""
    return os.path.join(os.path.dirname(os.path.abspath(__file__)), NOME_TEMPLATE)


def formatar_alternativas_markdown(alternativas):
    """Formata alternativas em Markdown para o preview."""
    letras = "abcdefghijklmnopqrstuvwxyz"
    partes = []
    for idx, alt in enumerate(alternativas):
        letra = letras[idx] if idx < len(letras) else str(idx + 1)
        partes.append(f"- **({letra})** {alt}")
    return "\n".join(partes)


def normalizar_alternativas(alternativas_texto):
    """Converte o texto livre das alternativas em uma lista limpa."""
    return [
        linha.strip()
        for linha in alternativas_texto.strip().split("\n")
        if linha.strip()
    ]


def inserir_questao(enunciado, alternativas_texto):
    """Adiciona uma questão ao session_state."""
    if not enunciado.strip():
        st.warning("Por favor, insira o enunciado da questão.")
        return
    alternativas = normalizar_alternativas(alternativas_texto)
    st.session_state.questoes.append(
        {"enunciado": enunciado.strip(), "alternativas": alternativas}
    )
    st.success(f"Questão {len(st.session_state.questoes)} inserida com sucesso!")


def limpar_questoes():
    """Remove todas as questões do session_state."""
    st.session_state.questoes = []
    st.info("Todas as questões foram removidas.")


def atualizar_questao(indice, enunciado, alternativas_texto):
    """Atualiza uma questão existente no session_state."""
    if not enunciado.strip():
        st.warning("Por favor, insira o enunciado da questão.")
        return False

    st.session_state.questoes[indice] = {
        "enunciado": enunciado.strip(),
        "alternativas": normalizar_alternativas(alternativas_texto),
    }
    st.success(f"Questão {indice + 1} atualizada com sucesso!")
    return True


@st.dialog("Confirmar exclusão")
def confirmar_limpeza_questoes():
    """Exibe um popup para confirmar a limpeza das questões."""
    st.warning("Deseja realmente excluir todas as questões adicionadas?")
    col_cancelar, col_confirmar = st.columns(2)

    with col_cancelar:
        if st.button("Cancelar", use_container_width=True):
            st.rerun()

    with col_confirmar:
        if st.button("Excluir tudo", type="primary", use_container_width=True):
            limpar_questoes()
            st.rerun()


@st.dialog("Editar questão")
def editar_questao_dialog(indice):
    """Abre um popup para editar uma questão existente."""
    questao = st.session_state.questoes[indice]
    alternativas_iniciais = "\n".join(questao["alternativas"])

    with st.form(key=f"form_editar_questao_{indice}"):
        enunciado = st.text_area(
            "Enunciado da Questão",
            value=questao["enunciado"],
            height=150,
        )
        alternativas_texto = st.text_area(
            "Alternativas (uma por linha, deixe em branco para questão discursiva)",
            value=alternativas_iniciais,
            height=150,
        )
        col_cancelar, col_salvar = st.columns(2)
        with col_cancelar:
            cancelar = st.form_submit_button("Cancelar", use_container_width=True)
        with col_salvar:
            salvar = st.form_submit_button(
                "Salvar alterações",
                type="primary",
                use_container_width=True,
            )

        if cancelar:
            st.rerun()

        if salvar and atualizar_questao(indice, enunciado, alternativas_texto):
            st.rerun()


def montar_nome_arquivo(componente, data_atividade, extensao="docx"):
    """Gera um nome de arquivo seguro usando componente curricular e data."""
    partes = ["Atividade"]

    if componente.strip():
        componente_limpo = re.sub(r'[<>:"/\\|?*]', "", componente.strip())
        componente_limpo = re.sub(r"\s+", " ", componente_limpo).strip(" .")
        if componente_limpo:
            partes.append(componente_limpo)
    else:
        partes.append("Componete")

    if data_atividade.strip():
        data_limpa = re.sub(r'[<>:"/\\|?*]', "-", data_atividade.strip())
        data_limpa = re.sub(r"\s+", " ", data_limpa).strip(" .")
        if data_limpa:
            partes.append(data_limpa)
    else:
        partes.append("data")

    nome_base = "-".join(partes)
    return f"{nome_base}.{extensao}"


def obter_data_para_widget(data_atividade):
    """Converte a data salva em texto para o formato aceito pelo date_input."""
    if not data_atividade.strip():
        return None

    try:
        return datetime.strptime(data_atividade, "%d/%m/%Y").date()
    except ValueError:
        return None


def _obter_texto_paragrafo(paragraph):
    """Retorna o texto completo de um parágrafo (concatena runs)."""
    return "".join(run.text for run in paragraph.runs)


def _substituir_marcador_em_paragrafo(paragraph, marcador, texto_substituicao, usar_negrito=False):
    """
    Tenta substituir um marcador dentro de um parágrafo preservando o máximo
    possível da formatação original. Retorna True se houve substituição.
    """
    texto_completo = _obter_texto_paragrafo(paragraph)
    if marcador not in texto_completo:
        return False

    novo_texto = texto_completo.replace(marcador, texto_substituicao)

    if len(paragraph.runs) == 1:
        run = paragraph.runs[0]
        run.text = novo_texto
        if usar_negrito:
            run.bold = True
        return True

    if len(paragraph.runs) > 1:
        run_referencia = paragraph.runs[0]
        fonte_nome = run_referencia.font.name
        fonte_tamanho = run_referencia.font.size
        cor = run_referencia.font.color.rgb if run_referencia.font.color and run_referencia.font.color.rgb else None
        negrito_original = run_referencia.bold
        italico = run_referencia.italic

        for run in paragraph.runs:
            run.text = ""

        novo_run = paragraph.runs[0]
        novo_run.text = novo_texto
        if fonte_nome:
            novo_run.font.name = fonte_nome
        if fonte_tamanho:
            novo_run.font.size = fonte_tamanho
        if cor:
            novo_run.font.color.rgb = cor
        if usar_negrito:
            novo_run.bold = True
        else:
            novo_run.bold = negrito_original
        novo_run.italic = italico
        return True

    return False


def _quebrar_linhas_no_run(run, texto):
    """
    Insere um texto no run, substituindo '\n' por quebras de linha reais do Word
    (mesmo parágrafo) através de add_break().
    """
    linhas = texto.split("\n")
    for i, linha in enumerate(linhas):
        run.add_text(linha)
        if i < len(linhas) - 1:
            run.add_break()


def _substituir_conteudo_questao_em_paragrafo(paragrafo, marcador, texto_subst):
    """Substitui o marcador de conteúdo da questão preservando a formatação base."""
    texto_novo = _obter_texto_paragrafo(paragrafo).replace(marcador, "")

    if len(paragrafo.runs) == 0:
        paragrafo.add_run(texto_novo)
        if texto_subst:
            run_questao = paragrafo.add_run()
            _quebrar_linhas_no_run(run_questao, texto_subst)
        return

    run_ref = paragrafo.runs[0]
    fonte_nome = run_ref.font.name
    fonte_tamanho = run_ref.font.size
    cor = (
        run_ref.font.color.rgb
        if run_ref.font.color and run_ref.font.color.rgb
        else None
    )
    negrito_ref = run_ref.bold

    for run in paragrafo.runs:
        run.text = ""

    paragrafo.runs[0].text = texto_novo
    if fonte_nome:
        paragrafo.runs[0].font.name = fonte_nome
    if fonte_tamanho:
        paragrafo.runs[0].font.size = fonte_tamanho
    if cor:
        paragrafo.runs[0].font.color.rgb = cor
    paragrafo.runs[0].bold = negrito_ref

    if texto_subst:
        run_conteudo = paragrafo.add_run()
        if fonte_nome:
            run_conteudo.font.name = fonte_nome
        if fonte_tamanho:
            run_conteudo.font.size = fonte_tamanho
        _quebrar_linhas_no_run(run_conteudo, texto_subst)


def _iterar_blocos(parent):
    """Itera parágrafos e tabelas na ordem em que aparecem no documento."""
    if isinstance(parent, DocxDocument):
        parent_element = parent.element.body
    elif isinstance(parent, _Cell):
        parent_element = parent._tc
    else:
        raise TypeError("Tipo de contêiner não suportado para iteração de blocos.")

    for child in parent_element.iterchildren():
        if child.tag.endswith("}p"):
            yield Paragraph(child, parent)
        elif child.tag.endswith("}tbl"):
            yield Table(child, parent)


def _obter_texto_bloco(bloco):
    """Retorna o texto de um parágrafo ou de uma tabela inteira."""
    if isinstance(bloco, Paragraph):
        return bloco.text
    if isinstance(bloco, Table):
        # A remoção por delimitadores acontece no nível do contêiner (documento ou
        # célula). Para tabelas, evitamos agregar o texto interno aqui para não
        # remover estruturas maiores do que o desejado no nível do documento.
        return ""
    return ""


def _remover_bloco(bloco):
    """Remove um parágrafo ou tabela do documento."""
    elemento = bloco._element
    elemento.getparent().remove(elemento)


def _iterar_celulas(container):
    """Retorna todas as células do documento, incluindo tabelas aninhadas."""
    tabelas = container.tables if hasattr(container, "tables") else []
    for tabela in tabelas:
        for row in tabela.rows:
            for cell in row.cells:
                yield cell
                yield from _iterar_celulas(cell)


def _iterar_containers(documento):
    """Itera o documento e todas as células onde um bloco de questão pode existir."""
    yield documento
    for cell in _iterar_celulas(documento):
        yield cell


def _encontrar_intervalo_bloco_generico(container):
    """Localiza o intervalo do bloco modelo único delimitado por INICIO/FIM."""
    marcador_inicio = "<<INICIO_QUESTAO>>"
    marcador_fim = "<<FIM_QUESTAO>>"
    blocos = list(_iterar_blocos(container))

    indice_inicio = None
    indice_fim = None

    for indice, bloco in enumerate(blocos):
        if not isinstance(bloco, Paragraph):
            continue
        texto_bloco = bloco.text.strip()
        if indice_inicio is None and marcador_inicio in texto_bloco:
            indice_inicio = indice
        if marcador_fim in texto_bloco:
            indice_fim = indice
            break

    if indice_inicio is None or indice_fim is None or indice_inicio > indice_fim:
        return None

    return blocos, indice_inicio, indice_fim


def _substituir_marcadores_genericos_em_blocos(blocos, numero_questao, questao):
    """Preenche um bloco clonado com título e conteúdo da questão."""
    titulo_questao = construir_titulo_questao(numero_questao)
    texto_questao = construir_texto_questao(questao)
    marcadores_titulo = {"<<TITULO_QUESTAO>>", "<<TituloQuestao>>"}
    marcadores_conteudo = {"<<QUESTAO>>", "<<Questao>>"}
    marcadores_controle = {"<<INICIO_QUESTAO>>", "<<FIM_QUESTAO>>"}

    for bloco in list(blocos):
        if isinstance(bloco, Paragraph):
            texto_limpo = bloco.text.strip()
            if texto_limpo in marcadores_controle:
                _substituir_marcador_em_paragrafo(bloco, texto_limpo, "")
                continue
            for marcador in marcadores_titulo:
                _substituir_marcador_em_paragrafo(bloco, marcador, titulo_questao)
            for marcador in marcadores_conteudo:
                if marcador in _obter_texto_paragrafo(bloco):
                    _substituir_conteudo_questao_em_paragrafo(
                        bloco, marcador, texto_questao
                    )
            for marcador in marcadores_controle:
                _substituir_marcador_em_paragrafo(bloco, marcador, "")
        elif isinstance(bloco, Table):
            for row in bloco.rows:
                for cell in row.cells:
                    for paragrafo in cell.paragraphs:
                        for marcador in marcadores_titulo:
                            _substituir_marcador_em_paragrafo(
                                paragrafo, marcador, titulo_questao
                            )
                        for marcador in marcadores_conteudo:
                            if marcador in _obter_texto_paragrafo(paragrafo):
                                _substituir_conteudo_questao_em_paragrafo(
                                    paragrafo, marcador, texto_questao
                                )
                        for marcador in marcadores_controle:
                            _substituir_marcador_em_paragrafo(paragrafo, marcador, "")


def _expandir_bloco_generico_questoes(documento, questoes):
    """
    Expande um único bloco modelo de questão para a quantidade de questões informada.
    O bloco modelo deve usar:
    <<INICIO_QUESTAO>>, <<TITULO_QUESTAO>>, <<QUESTAO>>, <<FIM_QUESTAO>>
    """
    for container in _iterar_containers(documento):
        intervalo = _encontrar_intervalo_bloco_generico(container)
        if intervalo is None:
            continue

        blocos, indice_inicio, indice_fim = intervalo
        blocos_modelo = blocos[indice_inicio : indice_fim + 1]
        xml_modelo = [deepcopy(bloco._element) for bloco in blocos_modelo]
        ancora = blocos_modelo[-1]._element

        for numero_questao, questao in enumerate(questoes, start=1):
            novos_blocos = []
            for xml in [deepcopy(elemento) for elemento in xml_modelo]:
                ancora.addnext(xml)
                ancora = xml
                if xml.tag.endswith("}p"):
                    novos_blocos.append(Paragraph(xml, container))
                elif xml.tag.endswith("}tbl"):
                    novos_blocos.append(Table(xml, container))
            _substituir_marcadores_genericos_em_blocos(
                novos_blocos, numero_questao, questao
            )

        for bloco in reversed(blocos_modelo):
            _remover_bloco(bloco)
        return True

    return False


def _remover_ou_limpar_bloco_questao_em_container(container, numero_questao, manter_conteudo):
    """
    Remove o bloco delimitado por <<INICIO_QUESTAO_N>> e <<FIM_QUESTAO_N>> quando
    a questão estiver vazia. Se a questão existir, remove apenas os marcadores.
    """
    marcador_inicio = f"<<INICIO_QUESTAO_{numero_questao}>>"
    marcador_fim = f"<<FIM_QUESTAO_{numero_questao}>>"
    blocos = list(_iterar_blocos(container))

    indice_inicio = None
    indice_fim = None

    for indice, bloco in enumerate(blocos):
        texto_bloco = _obter_texto_bloco(bloco)
        if indice_inicio is None and marcador_inicio in texto_bloco:
            indice_inicio = indice
        if marcador_fim in texto_bloco:
            indice_fim = indice
            break

    if indice_inicio is None or indice_fim is None or indice_inicio > indice_fim:
        return

    blocos_alvo = blocos[indice_inicio : indice_fim + 1]

    if manter_conteudo:
        for bloco in blocos_alvo:
            if isinstance(bloco, Paragraph):
                texto_limpo = bloco.text.strip()
                if texto_limpo in {marcador_inicio, marcador_fim}:
                    _remover_bloco(bloco)
                else:
                    _substituir_marcador_em_paragrafo(bloco, marcador_inicio, "")
                    _substituir_marcador_em_paragrafo(bloco, marcador_fim, "")
    else:
        for bloco in reversed(blocos_alvo):
            _remover_bloco(bloco)


def _remover_ou_limpar_bloco_questao(documento, numero_questao, manter_conteudo):
    """Aplica a remoção/limpeza do bloco da questão no documento e nas células."""
    _remover_ou_limpar_bloco_questao_em_container(
        documento, numero_questao, manter_conteudo
    )
    for cell in _iterar_celulas(documento):
        _remover_ou_limpar_bloco_questao_em_container(
            cell, numero_questao, manter_conteudo
        )


def construir_texto_questao(questao):
    """
    Monta o texto bruto de uma questão com enunciado, quebras de linha e
    alternativas formatadas com letras (a), (b), etc. Retorna string com '\n'.
    """
    partes = []
    partes.append(questao["enunciado"])
    if questao["alternativas"]:
        letras = "abcdefghijklmnopqrstuvwxyz"
        for idx, alt in enumerate(questao["alternativas"]):
            letra = letras[idx] if idx < len(letras) else str(idx + 1)
            partes.append(f"({letra}) {alt}")
    else:
        partes.append("")
        partes.append("")
    return "\n".join(partes)


def construir_titulo_questao(numero_questao):
    """Monta o título exibido no cabeçalho visual da questão."""
    return f"Questão {numero_questao:02d}"


def processar_substituicoes(
    documento, questoes, componente, data_atividade, professor, ano_serie
):
    """
    Realiza find-and-replace abrangente em paragraphs e table.cells.
    - Substitui <<Questao1>>, <<Questao2>>... pelas questões do usuário.
    - Marcadores não utilizados são substituídos por string vazia.
    - Substitui também marcadores de cabeçalho opcionais: <<Componente>>, <<Data>>, <<Professor>>
    """
    substituicoes = {}
    substituicoes["<<Componente>>"] = componente or ""
    substituicoes["<<Data>>"] = data_atividade or ""
    substituicoes["<<Professor>>"] = professor or ""
    substituicoes["<<ANO_SERIE>>"] = ano_serie or ""
    substituicoes["<<AnoSerie>>"] = ano_serie or ""
    for opcao_ano in OPCOES_ANO:
        substituicoes[f"<<{opcao_ano}>>"] = ano_serie or ""

    usa_bloco_generico = _expandir_bloco_generico_questoes(documento, questoes)

    if usa_bloco_generico:
        substituicoes["<<TITULO_QUESTAO>>"] = ""
        substituicoes["<<TituloQuestao>>"] = ""
        substituicoes["<<QUESTAO>>"] = ""
        substituicoes["<<Questao>>"] = ""
        substituicoes["<<INICIO_QUESTAO>>"] = ""
        substituicoes["<<FIM_QUESTAO>>"] = ""

    for n in range(1, 501):
        marcador = f"<<Questao{n}>>"
        marcador_titulo = f"<<TituloQuestao{n}>>"
        marcador_titulo_padrao = f"<<TITULO_QUESTAO_{n}>>"
        possui_questao = (n - 1) < len(questoes)
        if not usa_bloco_generico:
            _remover_ou_limpar_bloco_questao(documento, n, manter_conteudo=possui_questao)

        if (n - 1) < len(questoes):
            substituicoes[marcador] = construir_texto_questao(questoes[n - 1])
            titulo_questao = construir_titulo_questao(n)
            substituicoes[marcador_titulo] = titulo_questao
            substituicoes[marcador_titulo_padrao] = titulo_questao
        else:
            substituicoes[marcador] = ""
            substituicoes[marcador_titulo] = ""
            substituicoes[marcador_titulo_padrao] = ""

    todos_elementos = []

    def _coletar_paragrafos():
        for p in documento.paragraphs:
            todos_elementos.append(p)
        for table in documento.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        todos_elementos.append(p)
                    for inner_table in cell.tables:
                        for inner_row in inner_table.rows:
                            for inner_cell in inner_row.cells:
                                for p in inner_cell.paragraphs:
                                    todos_elementos.append(p)

    _coletar_paragrafos()

    for paragrafo in todos_elementos:
        texto_p = _obter_texto_paragrafo(paragrafo)
        if not texto_p:
            continue

        houve_match = False
        marcador_match = None
        texto_subst = None
        for marcador, subst in substituicoes.items():
            if marcador in texto_p:
                marcador_match = marcador
                texto_subst = subst
                houve_match = True
                break

        if not houve_match:
            continue

        if marcador_match.startswith("<<Questao"):
            _substituir_conteudo_questao_em_paragrafo(
                paragrafo, marcador_match, texto_subst
            )
        else:
            _substituir_marcador_em_paragrafo(paragrafo, marcador_match, texto_subst)


def gerar_documento_docx(questoes, componente, data_atividade, professor, ano_serie):
    """
    Abre o template, executa os replacements e retorna um BytesIO com o .docx
    final. Nunca altera o arquivo em disco.
    """
    caminho_template = obter_caminho_template()

    if not os.path.isfile(caminho_template):
        raise FileNotFoundError(
            f"Arquivo de template '{NOME_TEMPLATE}' não foi encontrado na raiz do projeto. "
            f"Certifique-se de que o arquivo exista em: {caminho_template}"
        )

    doc = Document(caminho_template)
    processar_substituicoes(
        doc, questoes, componente, data_atividade, professor, ano_serie
    )

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def converter_docx_para_pdf_bytes(docx_bytes):
    """
    Converte um DOCX em PDF usando o Microsoft Word via automação COM.
    Mantém o layout do template com mais fidelidade do que recriar o PDF.
    """
    try:
        import pythoncom
        from win32com.client import DispatchEx
    except ImportError as exc:
        raise RuntimeError(
            "A conversão para PDF requer pywin32 instalado e Microsoft Word disponível."
        ) from exc

    pythoncom.CoInitialize()
    word = None
    documento_word = None

    try:
        with tempfile.TemporaryDirectory() as pasta_temp:
            caminho_docx = os.path.join(pasta_temp, "atividade_temp.docx")
            caminho_pdf = os.path.join(pasta_temp, "atividade_temp.pdf")

            with open(caminho_docx, "wb") as arquivo_docx:
                arquivo_docx.write(docx_bytes.getvalue())

            word = DispatchEx("Word.Application")
            word.Visible = False
            word.DisplayAlerts = 0

            documento_word = word.Documents.Open(caminho_docx, ReadOnly=True)
            documento_word.SaveAs(caminho_pdf, FileFormat=17)
            documento_word.Close(False)
            documento_word = None
            word.Quit()
            word = None

            with open(caminho_pdf, "rb") as arquivo_pdf:
                return io.BytesIO(arquivo_pdf.read())
    except Exception as exc:
        raise RuntimeError(
            "Nao foi possivel converter o documento para PDF. "
            "Verifique se o Microsoft Word esta instalado e liberado para automacao."
        ) from exc
    finally:
        if documento_word is not None:
            documento_word.Close(False)
        if word is not None:
            word.Quit()
        pythoncom.CoUninitialize()


def obter_status_pdf():
    """Informa se a conversão para PDF está disponível no ambiente atual."""
    if os.name != "nt":
        return (
            False,
            "O PDF automatico fica disponivel apenas no uso local em Windows com Microsoft Word.",
        )

    try:
        import pythoncom  # noqa: F401
        from win32com.client import DispatchEx  # noqa: F401
    except ImportError:
        return (
            False,
            "O PDF automatico exige pywin32 instalado e Microsoft Word disponivel neste computador.",
        )

    return True, ""


def main():
    st.set_page_config(
        page_title="Gerador de Atividades - Template DOCX",
        page_icon="📝",
        layout="wide",
    )

    inicializar_estado()

    st.title("📝 Gerador de Listas de Exercícios (Template DOCX)")
    st.markdown("---")

    caminho_template = obter_caminho_template()
    template_existe = os.path.isfile(caminho_template)

    if not template_existe:
        st.error(
            f"❌ **Arquivo de template não encontrado:** `{NOME_TEMPLATE}`\n\n"
            f"Coloque o arquivo `.docx` modelo na seguinte pasta:\n"
            f"`{os.path.dirname(caminho_template)}`"
        )
        st.stop()

    with st.expander("⚙️  Configuração do Cabeçalho", expanded=True):
        col_cfg1, col_cfg2, col_cfg3, col_cfg4 = st.columns(4)
        with col_cfg1:
            componente = st.text_input(
                "Componente Curricular",
                value=st.session_state.componente,
                placeholder="Ex: Matemática, Português, Ciências...",
            )
            st.session_state.componente = componente
        with col_cfg2:
            data_selecionada = st.date_input(
                "Data",
                value=obter_data_para_widget(st.session_state.data_atividade),
                format="DD/MM/YYYY",
            )
            st.session_state.data_atividade = (
                data_selecionada.strftime("%d/%m/%Y") if data_selecionada else ""
            )
        with col_cfg3:
            professor = st.text_input(
                "Professor(a)",
                value=st.session_state.professor,
            )
            st.session_state.professor = professor
        with col_cfg4:
            ano_atual = (
                st.session_state.ano_serie
                if st.session_state.ano_serie in OPCOES_ANO
                else "4º ANO"
            )
            indice_ano = OPCOES_ANO.index(ano_atual)
            ano_serie = st.selectbox(
                "Ano / Série",
                options=OPCOES_ANO,
                index=indice_ano,
            )
            st.session_state.ano_serie = ano_serie

    col1, col2 = st.columns([1, 1], gap="large")

    with col1:
        st.subheader("➕ Inserir Questões")

        with st.form(key="form_questao", clear_on_submit=True):
            enunciado = st.text_area(
                "Enunciado da Questão",
                height=150,
                placeholder="Digite aqui o enunciado completo da questão...",
            )
            alternativas_texto = st.text_area(
                "Alternativas (uma por linha, deixe em branco para questão discursiva)",
                height=150,
                placeholder=(
                    "Digite uma alternativa por linha:\n"
                    "Alternativa A\n"
                    "Alternativa B\n"
                    "Alternativa C\n"
                    "\nDeixe em branco para questão discursiva."
                ),
            )
            submit = st.form_submit_button("Inserir na Atividade", type="primary")
            if submit:
                inserir_questao(enunciado, alternativas_texto)

        st.markdown("---")
        if st.button("🗑️  Limpar Todas as Questões", type="secondary"):
            confirmar_limpeza_questoes()

        st.markdown("---")
        st.subheader("📄 Gerar Documento")
        if not st.session_state.questoes:
            st.info("Adicione pelo menos uma questão para gerar o documento.")
        else:
            try:
                docx_bytes = gerar_documento_docx(
                    st.session_state.questoes,
                    st.session_state.componente,
                    st.session_state.data_atividade,
                    st.session_state.professor,
                    st.session_state.ano_serie,
                )
                nome_arquivo = montar_nome_arquivo(
                    st.session_state.componente,
                    st.session_state.data_atividade,
                )
                col_download_docx, col_download_pdf = st.columns(2)
                with col_download_docx:
                    st.download_button(
                        label="⬇️ Baixar Atividade (DOCX)",
                        data=docx_bytes,
                        file_name=nome_arquivo,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        type="primary",
                        use_container_width=True,
                    )
                with col_download_pdf:
                    pdf_disponivel, mensagem_pdf = obter_status_pdf()
                    if pdf_disponivel:
                        pdf_bytes = converter_docx_para_pdf_bytes(docx_bytes)
                        nome_pdf = montar_nome_arquivo(
                            st.session_state.componente,
                            st.session_state.data_atividade,
                            extensao="pdf",
                        )
                        st.download_button(
                            label="⬇️ Baixar Atividade (PDF)",
                            data=pdf_bytes,
                            file_name=nome_pdf,
                            mime="application/pdf",
                            use_container_width=True,
                        )
                    else:
                        st.button(
                            "⬇️ Baixar Atividade (PDF)",
                            disabled=True,
                            use_container_width=True,
                        )
                        st.caption(mensagem_pdf)
                st.success(
                    f"Pronto! {len(st.session_state.questoes)} questão(ões) inserida(s) no template."
                )
            except FileNotFoundError as e:
                st.error(str(e))
            except Exception as e:
                st.error(f"Erro ao gerar o documento: {e}")

    with col2:
        st.subheader("👀 Preview das Questões")
        if not st.session_state.questoes:
            st.info(
                "Nenhuma questão adicionada ainda. "
                "Use o formulário ao lado para inserir questões no template."
            )
        else:
            for idx, questao in enumerate(st.session_state.questoes, start=1):
                with st.container(border=True):
                    col_titulo, col_acao = st.columns([4, 1])
                    with col_titulo:
                        st.markdown(f"### Questão {str(idx).zfill(2)}")
                    with col_acao:
                        if st.button("Editar", key=f"editar_questao_{idx}"):
                            editar_questao_dialog(idx - 1)
                    st.markdown(f"**{questao['enunciado']}**")
                    if questao["alternativas"]:
                        st.markdown(formatar_alternativas_markdown(questao["alternativas"]))
                    else:
                        st.caption("💬 Questão discursiva — espaço para resposta")

            st.markdown("---")
            st.caption(f"Total de questões: **{len(st.session_state.questoes)}**")


if __name__ == "__main__":
    main()
